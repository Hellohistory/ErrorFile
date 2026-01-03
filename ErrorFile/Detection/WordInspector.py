# ErrorFile/Detection/WordInspector.py

import zipfile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from ..report import (
    TAG_CORRUPTED,
    TAG_INVALID_FORMAT,
    TAG_IO_ERROR,
    fail_finding,
    ok_finding,
)


def _fast_check_docx(file_path):
    try:
        with zipfile.ZipFile(file_path, "r") as archive:
            names = set(archive.namelist())
            required = {"[Content_Types].xml", "word/document.xml"}
            if not required.issubset(names):
                return fail_finding(
                    "DOCX missing required parts; file may be corrupted.",
                    TAG_CORRUPTED,
                    TAG_INVALID_FORMAT,
                )
    except zipfile.BadZipFile as exc:
        return fail_finding(
            f"DOCX is not a valid ZIP container: {exc}",
            TAG_INVALID_FORMAT,
            error=str(exc),
        )
    except Exception as exc:
        return fail_finding(
            f"DOCX fast check failed: {exc}",
            TAG_IO_ERROR,
            error=str(exc),
        )
    return ok_finding("DOCX fast check passed.")


def check_docx_file(file_path, mode="deep"):
    """Check .docx integrity."""
    if mode == "fast":
        return _fast_check_docx(file_path)
    try:
        doc = Document(file_path)

        for paragraph in doc.paragraphs:
            _ = paragraph.text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _ = cell.text

        return ok_finding("DOCX deep check passed.")
    except PackageNotFoundError as exc:
        return fail_finding(
            f"DOCX corrupted or invalid: {exc}",
            TAG_CORRUPTED,
            TAG_INVALID_FORMAT,
            error=str(exc),
        )
    except Exception as exc:
        return fail_finding(
            f"DOCX deep check failed: {exc}",
            TAG_IO_ERROR,
            error=str(exc),
        )
